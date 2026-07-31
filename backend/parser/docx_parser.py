import time
import os
import re
from typing import Dict, List, Any, Optional
import docx

class FastDOCXParser:
    """
    Production-Grade, High-Speed DOCX Parser using python-docx.
    Parses OpenXML DOM structure into Headings, Paragraphs, Lists, Tables, and Hyperlinks.
    Runs in < 0.1s for 50-page documents.
    """
    def parse(self, docx_path: str) -> Dict[str, Any]:
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"DOCX file not found at: {docx_path}")

        start_time = time.perf_counter()
        doc = docx.Document(docx_path)
        doc_name = os.path.basename(docx_path)

        elements = []
        tables = []

        # 1. Parse Elements (Paragraphs, Headings, Lists) in document body order
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue

            style_name = p.style.name.lower() if p.style else ""
            
            # Extract hyperlinks if embedded in paragraph XML
            text_with_links = self._extract_paragraph_links(p)
            if text_with_links:
                text = text_with_links

            if "title" in style_name or "heading 1" in style_name or style_name == "h1":
                element_type = "H1_TITLE"
            elif "heading 2" in style_name or style_name == "h2":
                element_type = "H2_HEADING"
            elif "heading 3" in style_name or style_name == "h3":
                element_type = "H3_SUBHEADING"
            elif "list" in style_name or text.startswith(("•", "-", "*", "1.", "2.", "3.")):
                element_type = "LIST_ITEM"
            else:
                element_type = "PARAGRAPH"

            elements.append({
                "type": element_type,
                "text": text,
                "style": style_name
            })

        # 2. Parse Tables
        for tab_idx, table in enumerate(doc.tables):
            extracted_grid = []
            for row in table.rows:
                row_cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
                if any(row_cells):
                    extracted_grid.append(row_cells)

            if len(extracted_grid) >= 2 and len(extracted_grid[0]) >= 1:
                tables.append({
                    "table_id": tab_idx + 1,
                    "rows": extracted_grid
                })

        elapsed_time = time.perf_counter() - start_time

        # DOCX documents are treated as page 1 (or continuous stream) with page metadata
        return {
            "doc_name": doc_name,
            "total_pages": 1,
            "processing_time_seconds": round(elapsed_time, 4),
            "pages_per_second": round(1 / elapsed_time, 2) if elapsed_time > 0 else 0,
            "pages": [{
                "page": 1,
                "elements": elements,
                "tables": tables
            }]
        }

    def _extract_paragraph_links(self, paragraph) -> Optional[str]:
        """Extracts hyperlinks from paragraph XML elements."""
        try:
            has_link = False
            full_text = ""
            for child in paragraph._element:
                if child.tag.endswith('hyperlink'):
                    has_link = True
                    rId = child.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                    link_text = "".join(node.text for node in child.iter() if node.text)
                    if rId and link_text:
                        rel = paragraph.part.rels.get(rId)
                        if rel and rel.target_ref:
                            full_text += f"[{link_text.strip()}]({rel.target_ref}) "
                        else:
                            full_text += link_text + " "
                elif child.tag.endswith('r'):
                    text = "".join(node.text for node in child.iter() if node.text)
                    full_text += text
            return full_text.strip() if has_link and full_text else None
        except Exception:
            return None
